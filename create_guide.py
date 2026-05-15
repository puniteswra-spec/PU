from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

title = doc.add_heading('Remote Monitor System - Complete Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1a, 0x5e, 0x20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Monitor System designed by Puneet Upreti')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
today_str = datetime.date.today().strftime('%B %d, %Y')
run = p.add_run('Version 7.0.0 | Generated: ' + today_str)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. System Overview',
    '2. Architecture',
    '3. Components',
    '   3.1 Agent (SystemHelper.exe)',
    '   3.2 Server (server.source.js)',
    '   3.3 Dashboard (dashboard.html)',
    '   3.4 View Page (/view/:agentId)',
    '   3.5 Multi-Control Panel (/multi-control)',
    '   3.6 Support Session (/support/:token)',
    '4. Installation and Setup',
    '   4.1 Server Setup (Render.com)',
    '   4.2 Server Setup (Local PC)',
    '   4.3 Server Setup (Linux VPS)',
    '   4.4 Agent Deployment',
    '5. Features and Usage',
    '6. Configuration Files',
    '7. Security and Authentication',
    '8. Troubleshooting',
    '9. API Reference',
    '10. File Structure'
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith('   '):
        for run in p.runs:
            run.bold = True

doc.add_page_break()

doc.add_heading('1. System Overview', level=1)
doc.add_paragraph('The Remote Monitor System is a zero-config, self-healing, multi-agent remote monitoring and control platform. It allows a central administrator to:')
items = [
    'View live screens of multiple remote computers simultaneously (CCTV-style grid)',
    'Take remote control of any agent (mouse, keyboard)',
    'Control multiple agents simultaneously from a single panel',
    'Generate shareable support links with time-limited access',
    'Push software updates to all agents remotely',
    'Monitor system activity, idle/active time, uptime/downtime',
    'Export logs to CSV and push to GitHub',
    'Expose agents via tunnel for global access'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Key Design Principles:')
run.bold = True
run.font.size = Pt(12)

principles = [
    'Zero-config: Agent works out of the box with embedded server URLs',
    'Self-healing: Auto-reconnects with smart backoff, watchdog restart',
    'Multi-server fallback: Tries multiple servers in priority order',
    'Secure: SHA256 token auth, password-protected control, session expiry',
    'Scalable: Supports 100+ agents with dynamic grid layout'
]
for p_text in principles:
    doc.add_paragraph(p_text, style='List Bullet')

doc.add_page_break()

doc.add_heading('2. Architecture', level=1)
doc.add_paragraph('The system uses a hub-and-spoke architecture with dual connections for optimal performance:')

doc.add_heading('Connection Flow', level=2)
doc.add_paragraph('Each agent maintains TWO WebSocket connections:')

table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Connection', 'Purpose', 'URL']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

data = [
    ['PRIMARY (Cloud)', 'Control + Registration + Frames', 'wss://pu-k752.onrender.com'],
    ['SECONDARY (Local)', 'Frames only (low latency)', 'ws://127.0.0.1:3000'],
    ['FALLBACK (Tunnel)', 'Auto-started if cloud fails 3x', 'bore.pub / localhost.run']
]
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        table.rows[r].cells[c].text = val
        for paragraph in table.rows[r].cells[c].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Why dual connections? The primary handles all control commands and registration. The secondary sends frames to the local server for ultra-low-latency viewing (<10ms) when the dashboard is on the same network. The secondary does NOT register as an agent, preventing race conditions.')

doc.add_page_break()

doc.add_heading('3. Components', level=1)

doc.add_heading('3.1 Agent (SystemHelper.exe)', level=2)
doc.add_paragraph('Language: Go (compiled to Windows executable)')
doc.add_paragraph('Version: 7.0.0')
doc.add_paragraph('Source: main.go, exec_windows.go, exec_darwin.go, exec_other.go')

doc.add_heading('Startup Sequence:', level=3)
steps = [
    'init() - Loads hostname, auth credentials, custom URLs from urls.ini, fetches GitHub registry',
    'main() - Signal handler, prevents duplicate instances, loads agent ID',
    'setupAutostart() - Creates watchdog.vbs in %APPDATA%\\SystemHelper, adds to registry Run key',
    'startConfigServer() - Starts local config web server on port 8181',
    'Server discovery - Checks localhost:3000, scans network, may start own server mode',
    'connect() - Dials URLs in order: Render -> Local -> Direct IP',
    'Sends agent-hello with connectionId (nanosecond timestamp) for race prevention',
    'Starts secondary goroutine for local frame forwarding (no registration)',
    'Main loop: reads PRIMARY WS for commands, captures frames, sends to both connections'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(str(i) + '. ' + step)

doc.add_heading('Command Handlers (Primary WS):', level=3)
cmds = [
    'set-fps - Adjust frame capture rate (1-30 fps)',
    'control - Execute mouse/keyboard commands (rate limited: 30/sec)',
    'push-update - Download and apply .exe update, restart',
    'switch-server - Change primary server URL, save to urls.ini',
    'file-transfer - Receive and save file from dashboard',
    'request-file - Read and send file to dashboard',
    'start-tunnel - Start bore.pub/SSH tunnel for global access',
    'cleanup-logs - Delete old logs, truncate current',
    'become-server - Designate this agent as fallback server',
    'webrtc-offer/ice-candidate - WebRTC signaling for P2P streaming'
]
for cmd in cmds:
    doc.add_paragraph(cmd, style='List Bullet')

doc.add_heading('Activity Tracking:', level=3)
doc.add_paragraph('The agent tracks user activity using GetLastInputInfo (Windows API):')
doc.add_paragraph('- Idle threshold: 5 minutes of no input', style='List Bullet')
doc.add_paragraph('- Logs: system-wake, system-sleep, status-update events', style='List Bullet')
doc.add_paragraph('- CSV format: Date, Time, Agent ID, Hostname, Local IP, Public IP, Event, Details, Uptime, Idle, Active, Downtime, State', style='List Bullet')

doc.add_heading('3.2 Server (server.source.js)', level=2)
doc.add_paragraph('Runtime: Node.js 18+ on Render.com or any server')
doc.add_paragraph('Dependencies: express, ws')
doc.add_paragraph('Entry: package.json -> "start": "node server.source.js"')

doc.add_heading('HTTP Routes:', level=3)
routes = [
    'GET / - Serves dashboard.html with TOKEN_PLACEHOLDER replaced',
    'GET /view/:agentId - Full-screen single agent viewer (password for control)',
    'GET /multi-control - Unified multi-agent control panel',
    'GET /support/:token - Time-limited support session with temp access',
    'GET /remote-assistant - Browser-based screen sharing (no install, works on mobile/desktop)',
    'POST /api/support-token - Generate support session token (admin only)',
    'POST /api/upload-update - Push .exe to all agents via WS',
    'POST /api/send-file/:agentId - Send file to specific agent',
    'POST /api/switch-server - Tell agents to switch server URL',
    'POST /api/tunnel/:agentId - Start tunnel on agent',
    'GET /api/report - Agent report (JSON/CSV/HTML)',
    'POST /api/cleanup - Clear server history, notify agents',
    'GET /api/agents - List connected agents',
    'GET /api/frame/:agentId - Get latest frame',
    'GET /api/logs/:agentId? - Get agent logs',
    'GET /api/export-logs - Download CSV logs',
    'GET /api/compile-monthly-report - Compile monthly JSON report',
    'POST /api/push-logs-to-github - Commit logs to GitHub repo'
]
for route in routes:
    doc.add_paragraph(route, style='List Bullet')

doc.add_heading('WebSocket Events (/ws):', level=3)
doc.add_paragraph('Authentication: ?token=SHA256(username:password) or Basic Auth header')
ws_events = [
    'agent-hello - Register agent (with connectionId for race prevention)',
    'agent-frame - Store frame, broadcast to all viewers',
    'agent-status - Update stats, log, broadcast to dashboards',
    'dashboard-hello - Register dashboard, send agent list, CCTV mode',
    'view-agent - Add viewer to agent, send frame, increase FPS',
    'stop-viewing - Remove viewer, decrease FPS if no viewers',
    'support-view - Support session viewer (token-validated)',
    'support-control - Toggle control in support session',
    'control - Forward control command to agent (dashboard or support)',
    'request-file - Forward file request to agent',
    'file-response - Broadcast file from agent to dashboards',
    'file-transfer - Forward file from dashboard to agent',
    'tunnel-status - Broadcast tunnel URL from agent',
    'push-update - Push update to all agents',
    'become-server - Forward to agent',
    'webrtc-offer/answer/ice-candidate - WebRTC signaling relay'
]
for ev in ws_events:
    doc.add_paragraph(ev, style='List Bullet')

doc.add_heading('3.3 Dashboard (dashboard.html)', level=2)
doc.add_paragraph('A white-themed CCTV-style grid dashboard with:')
features = [
    'Dynamic grid layout (1-6 columns based on agent count)',
    'Live screen thumbnails with name, IP, hostname',
    'No auth popup - direct access, lock icon for admin features',
    'Admin features (password: puneet12): Push Update, Remote Control, Show/Hide, Export Logs, Push to GitHub',
    'Remote Control dropdown: checkboxes to select agents, Take Control or Multi-Control',
    'Hide/Show dropdown: checkboxes per agent, persisted in localStorage',
    'Per-agent buttons: Share Link, View, Expose (tunnel), Support Link',
    'Tunnel URL display per agent (shown when tunnel starts)',
    'Wizard: switch all agents to Render/Local/VPS server'
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('3.4 View Page (/view/:agentId)', level=2)
doc.add_paragraph('Full-screen single agent viewer:')
doc.add_paragraph('- View-only by default (no auth required)', style='List Bullet')
doc.add_paragraph('- Click "Request Control" -> password prompt -> enable control', style='List Bullet')
doc.add_paragraph('- Mouse move, click, right-click, keyboard forwarding', style='List Bullet')
doc.add_paragraph('- Auto-hide UI after 4s, show on mouse move', style='List Bullet')
doc.add_paragraph('- FPS counter, connection status, fullscreen button', style='List Bullet')

doc.add_heading('3.5 Multi-Control Panel (/multi-control)', level=2)
doc.add_paragraph('Unified control for multiple agents in a single browser tab:')
doc.add_paragraph('- Grid layout of selected agents', style='List Bullet')
doc.add_paragraph('- Each screen has independent control toggle (OFF/ON)', style='List Bullet')
doc.add_paragraph('- Password required per screen to enable control', style='List Bullet')
doc.add_paragraph('- Mouse and keyboard sent to correct agent based on cursor position', style='List Bullet')

doc.add_heading('3.6 Support Session (/support/:token)', level=2)
doc.add_paragraph('Shareable, time-limited remote assistance link:')
doc.add_paragraph('- Generated by admin via Support button on agent cell', style='List Bullet')
doc.add_paragraph('- Configurable expiry (default 60 minutes)', style='List Bullet')
doc.add_paragraph('- View-only by default, password for control', style='List Bullet')
doc.add_paragraph('- Auto-cleanup when session expires', style='List Bullet')
doc.add_paragraph('- Shows countdown timer in header', style='List Bullet')

doc.add_heading('3.7 Remote Assistant (/remote-assistant)', level=2)
doc.add_paragraph('Browser-based screen sharing — no installation required. Works on any device:')
doc.add_paragraph('- Desktop (Windows/Mac/Linux): Share screen, window, or browser tab', style='List Bullet')
doc.add_paragraph('- Mobile (iOS/Android): Share camera view (screen sharing not supported on mobile browsers)', style='List Bullet')
doc.add_paragraph('- Appears as a tile in the dashboard grid like any other agent', style='List Bullet')
doc.add_paragraph('- Access via the "Remote Assistant" link in the dashboard topbar', style='List Bullet')
doc.add_paragraph('- Share the link: https://your-server.com/remote-assistant', style='List Bullet')
doc.add_paragraph('- The remote user clicks "Start Sharing", selects their screen, and you can see it live', style='List Bullet')
doc.add_paragraph('- View-only (no remote control for browser sessions)', style='List Bullet')
doc.add_paragraph('- Auto-reconnects if connection drops', style='List Bullet')

doc.add_page_break()

doc.add_heading('4. Installation and Setup', level=1)

doc.add_heading('4.1 Server Setup (Render.com)', level=2)
steps_render = [
    'Create a Web Service on Render.com',
    'Connect your GitHub repository (puniteswra-spec/PU)',
    'Build Command: npm install',
    'Start Command: node server.source.js',
    'Environment: Node 18.x',
    'Deploy - dashboard available at https://your-app.onrender.com',
    'Update GitHub registry (urls.ini) with the new wss:// URL'
]
for i, s in enumerate(steps_render, 1):
    doc.add_paragraph(str(i) + '. ' + s)

doc.add_heading('4.2 Server Setup (Local PC)', level=2)
doc.add_paragraph('Requirements: Windows PC, Node.js installed')
steps_local = [
    'Install Node.js from https://nodejs.org (LTS version)',
    'Create folder: C:\\RemoteMonitor\\',
    'Copy files: server.source.js, package.json, dashboard.html',
    'Open PowerShell in folder: npm install',
    'Start server: node server.source.js',
    'Dashboard: http://<PC-IP>:3000',
    'Auto-start: Create shortcut to start-server.bat in shell:startup'
]
for i, s in enumerate(steps_local, 1):
    doc.add_paragraph(str(i) + '. ' + s)

doc.add_heading('4.3 Server Setup (Linux VPS)', level=2)
doc.add_paragraph('Requirements: Ubuntu/Debian VPS, $5-10/month')
p = doc.add_paragraph()
run = p.add_run('curl -fsSL https://deb.nodesource.com/setup_18.x | sudo -E bash -\nsudo apt-get install -y nodejs\nmkdir -p /opt/remotemonitor\ncd /opt/remotemonitor\nnpm install --production\nsudo npm install -g pm2\npm2 start server.source.js --name remotemonitor\npm2 save\npm2 startup\nsudo ufw allow 3000/tcp')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('4.4 Agent Deployment', level=2)
steps_agent = [
    'Compile: go build -o SystemHelper.exe -ldflags="-s -w -H=windowsgui" .',
    'Distribute SystemHelper.exe to target machines',
    'Double-click to run (runs hidden in background)',
    'Agent auto-registers with server, watchdog ensures persistence',
    'Config panel: http://localhost:8181 (local only)'
]
for i, s in enumerate(steps_agent, 1):
    doc.add_paragraph(str(i) + '. ' + s)

doc.add_page_break()

doc.add_heading('5. Features and Usage', level=1)

doc.add_heading('5.1 Dashboard Overview', level=2)
doc.add_paragraph('The dashboard shows a grid of all connected agents. Each cell displays:')
doc.add_paragraph('- Agent name and LIVE indicator', style='List Bullet')
doc.add_paragraph('- Local IP, Public IP, Hostname', style='List Bullet')
doc.add_paragraph('- Tunnel URL (when active)', style='List Bullet')
doc.add_paragraph('- Action buttons: Share, View, Expose, Support', style='List Bullet')

doc.add_heading('5.2 Admin Login', level=2)
doc.add_paragraph('Click the lock icon in the top-right corner. Enter password: puneet12')
doc.add_paragraph('Unlocks: Push Update, Remote Control dropdown, Show/Hide dropdown, Export Logs, Compile, Push to GitHub, Expose button, Support Link button')

doc.add_heading('5.3 Remote Control', level=2)
steps_rc = [
    'Login as admin',
    'Click Remote Control dropdown',
    'Check agents you want to control',
    'Click "Take Control" - opens each in new tab',
    'Enter password in each tab to enable control'
]
for i, s in enumerate(steps_rc, 1):
    doc.add_paragraph(str(i) + '. ' + s)

doc.add_heading('5.4 Multi-Agent Control', level=2)
steps_mc = [
    'Login as admin',
    'Click Remote Control dropdown',
    'Check agents',
    'Click "Multi-Control" - opens single tab with grid',
    'Click OFF on any screen -> enter password -> control enabled'
]
for i, s in enumerate(steps_mc, 1):
    doc.add_paragraph(str(i) + '. ' + s)

doc.add_heading('5.5 Shareable Support Links', level=2)
steps_sl = [
    'Login as admin',
    'Click Support button on agent cell',
    'Enter expiry time (default 60 min)',
    'Link copied to clipboard',
    'Share link - recipient can view immediately, password for control'
]
for i, s in enumerate(steps_sl, 1):
    doc.add_paragraph(str(i) + '. ' + s)

doc.add_heading('5.6 Tunnel/Expose Agent', level=2)
steps_tunnel = [
    'Login as admin',
    'Click Expose (plug) button on agent cell',
    'Agent starts bore.pub tunnel (or SSH fallback)',
    'Tunnel URL appears in agent cell',
    'Use URL for global access to that agent'
]
for i, s in enumerate(steps_tunnel, 1):
    doc.add_paragraph(str(i) + '. ' + s)

doc.add_heading('5.7 Push Updates', level=2)
steps_update = [
    'Login as admin',
    'Click Push Update',
    'Select .exe file',
    'Update pushed to all connected agents',
    'Agents auto-restart with new version'
]
for i, s in enumerate(steps_update, 1):
    doc.add_paragraph(str(i) + '. ' + s)

doc.add_heading('5.8 Export Logs', level=2)
steps_logs = [
    'Login as admin',
    'Click Export Logs - downloads CSV',
    'Click Compile - generates monthly report',
    'Click Push to GitHub - commits logs to repo'
]
for i, s in enumerate(steps_logs, 1):
    doc.add_paragraph(str(i) + '. ' + s)

doc.add_page_break()

doc.add_heading('6. Configuration Files', level=1)

table2 = doc.add_table(rows=6, cols=3)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
headers2 = ['File', 'Location', 'Purpose']
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

config_data = [
    ['auth.ini', '%APPDATA%\\SystemHelper\\', 'Username and password'],
    ['urls.ini', '%APPDATA%\\SystemHelper\\ or next to .exe', 'Custom server URLs (priority order)'],
    ['agent.id', '%APPDATA%\\SystemHelper\\', 'Persistent agent identifier'],
    ['agent.ini', '%APPDATA%\\SystemHelper\\', 'Server preference flag'],
    ['tunnel.url', '%APPDATA%\\SystemHelper\\', 'Current tunnel URL']
]
for r, row_data in enumerate(config_data, 1):
    for c, val in enumerate(row_data):
        table2.rows[r].cells[c].text = val
        for paragraph in table2.rows[r].cells[c].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_page_break()

doc.add_heading('7. Security and Authentication', level=1)
doc.add_paragraph('Authentication Model:')
doc.add_paragraph('- Dashboard: No auth required for viewing. Lock icon unlocks admin features.', style='List Bullet')
doc.add_paragraph('- View pages: No auth required. Password needed for control.', style='List Bullet')
doc.add_paragraph('- Support sessions: Time-limited token. Password needed for control.', style='List Bullet')
doc.add_paragraph('- WebSocket: SHA256 token or Basic Auth required.', style='List Bullet')
doc.add_paragraph('- API endpoints: Basic Auth required for admin operations.', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Default Credentials:')
run.bold = True
doc.add_paragraph('Username: puneet')
doc.add_paragraph('Password: puneet12')
doc.add_paragraph('Change in: auth.ini or server.source.js (AUTH_USER/AUTH_PASS)')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Race Condition Prevention:')
run.bold = True
doc.add_paragraph('Each agent connection includes a nanosecond-precision connectionId. The server rejects stale connections (older connectionId) to prevent duplicate agent entries during rapid reconnection.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Rate Limiting:')
run.bold = True
doc.add_paragraph('Control commands are limited to 30 per second per agent. Excess commands are silently dropped to prevent flooding.')

doc.add_page_break()

doc.add_heading('8. Troubleshooting', level=1)

issues = [
    ('Agent not connecting', 'Check urls.ini, verify server is running, check firewall rules, check agent.log'),
    ('Remote control not working', 'Verify primary connection is active (not secondary), check rate limiting, ensure password entered'),
    ('Tunnel not starting', 'Check SSH availability, verify bore.exe downloaded, check network connectivity'),
    ('Dashboard shows blank screens', 'Check WebSocket connection, verify agent is sending frames, check browser console'),
    ('Agent keeps reconnecting', 'Check server stability, verify URLs in urls.ini, check network latency'),
    ('Logs not exporting', 'Check server has write permissions to logs/ directory'),
    ('Push update fails', 'Verify .exe file is valid, check agent has write permissions to its directory')
]
for title_text, desc in issues:
    p = doc.add_paragraph()
    run = p.add_run(title_text + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

doc.add_heading('9. API Reference', level=1)
doc.add_paragraph('All API endpoints require Basic Auth (puneet:puneet12) unless noted.')

api_table = doc.add_table(rows=18, cols=3)
api_table.style = 'Light Grid Accent 1'
api_table.alignment = WD_TABLE_ALIGNMENT.CENTER
api_headers = ['Endpoint', 'Method', 'Description']
for i, h in enumerate(api_headers):
    cell = api_table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

api_data = [
    ['/api/upload-update', 'POST', 'Push .exe to all agents'],
    ['/api/send-file/:agentId', 'POST', 'Send file to agent'],
    ['/api/switch-server', 'POST', 'Switch agent server URL'],
    ['/api/make-server/:agentId', 'POST', 'Make agent a server'],
    ['/api/tunnel/:agentId', 'POST', 'Start tunnel on agent'],
    ['/api/report', 'GET', 'Agent report (format=json/csv/html)'],
    ['/api/cleanup', 'POST', 'Clear logs and history'],
    ['/api/agents', 'GET', 'List connected agents'],
    ['/api/frame/:agentId', 'GET', 'Get latest frame'],
    ['/api/logs/:agentId?', 'GET', 'Get agent logs'],
    ['/api/export-logs', 'GET', 'Download CSV logs'],
    ['/api/compile-monthly-report', 'GET', 'Compile monthly report'],
    ['/api/push-logs-to-github', 'POST', 'Push logs to GitHub'],
    ['/api/support-token', 'POST', 'Generate support session token'],
    ['/api/status', 'GET', 'Agent config status (port 8181)'],
    ['/api/urls', 'GET/POST', 'Manage server URLs (port 8181)'],
    ['/api/restart', 'POST', 'Restart agent (port 8181)']
]
for r, row_data in enumerate(api_data, 1):
    for c, val in enumerate(row_data):
        api_table.rows[r].cells[c].text = val
        for paragraph in api_table.rows[r].cells[c].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_page_break()

doc.add_heading('10. File Structure', level=1)

file_tree = '''RemoteMonitor-Merged/
  main.go                    - Agent main logic (Go)
  exec_windows.go            - Windows-specific: autostart, watchdog, control
  exec_darwin.go             - macOS-specific stubs
  exec_other.go              - Linux/other stubs
  go.mod / go.sum            - Go dependencies
  SystemHelper.exe           - Compiled agent binary
  server.source.js           - Node.js server (Render deployment)
  dashboard.html             - White CCTV dashboard
  package.json               - Node.js dependencies
  urls.ini                   - Server URL registry
  config.json                - Agent config (optional)
  logs/                      - Server-side log storage
    agent-logs-YYYY-MM.csv
    report-YYYY-MM.json
  watchdog.vbs               - Auto-generated watchdog script
  RemoteMonitor.bat          - Windows batch launcher
  README.md                  - Project readme'''

p = doc.add_paragraph()
run = p.add_run(file_tree)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.save('P:/Opencode/RemoteMonitor-Merged/Remote_Monitor_System_Complete_Guide.docx')
print('Complete guide saved successfully')
