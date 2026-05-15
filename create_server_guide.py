from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

title = doc.add_heading('Server Change Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1a, 0x5e, 0x20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('How to Change Servers - Complete Reference')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
today_str = datetime.date.today().strftime('%B %d, %Y')
run = p.add_run('Version 7.0.0 | Monitor System designed by Puneet Upreti')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

doc.add_heading('Overview', level=1)
doc.add_paragraph('This guide explains how to change the server that agents connect to. The system supports multiple server types and provides several methods to switch servers without modifying code.')

doc.add_heading('Server URL Priority Order', level=1)
doc.add_paragraph('Agents try URLs in this order (first successful connection wins):')

table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Priority', 'Source', 'Example']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

data = [
    ['1', 'GitHub Registry (urls.ini in repo)', 'wss://new-server.onrender.com'],
    ['2', 'Local urls.ini (next to .exe)', 'ws://192.168.1.100:3000'],
    ['3', 'Local urls.ini (%APPDATA%)', 'ws://10.0.0.5:3000'],
    ['4', 'Default: Render.com', 'wss://pu-k752.onrender.com'],
    ['5', 'Local server', 'ws://127.0.0.1:3000'],
    ['6', 'Direct IP', 'ws://43.247.40.101:3000']
]
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        table.rows[r].cells[c].text = val
        for paragraph in table.rows[r].cells[c].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Important: The FIRST URL in the list that successfully connects becomes the PRIMARY connection. All control commands go through this connection.')

doc.add_page_break()

doc.add_heading('Server Types', level=1)

doc.add_heading('Type 1: Local Computer as Server', level=2)
doc.add_paragraph('Best for: Office LAN, same building, no internet required')
doc.add_paragraph('Cost: Free (uses existing PC)')
doc.add_paragraph('Access: LAN only (computers on same network)')

p = doc.add_paragraph()
run = p.add_run('Setup Steps:')
run.bold = True

steps = [
    'Install Node.js on the server PC (https://nodejs.org - LTS version)',
    'Create a folder: C:\\RemoteMonitor\\',
    'Copy these files to the folder: server.source.js, package.json, dashboard.html',
    'Open PowerShell in that folder and run: npm install',
    'Start the server: node server.source.js',
    'Find the PC IP address: ipconfig (look for IPv4 Address)',
    'Dashboard is now available at: http://<PC-IP>:3000',
    'Configure agents: Edit urls.ini to add ws://<PC-IP>:3000 as first URL'
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(str(i) + '. ' + s)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Make it auto-start on boot:')
run.bold = True

doc.add_paragraph('Create a file called start-server.bat in the folder:')
p = doc.add_paragraph()
run = p.add_run('@echo off\ncd /d C:\\RemoteMonitor\nnode server.source.js')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph('Place a shortcut to this file in the Startup folder:')
doc.add_paragraph('Press Win+R, type: shell:startup, press Enter')
doc.add_paragraph('Right-click in the folder -> New -> Shortcut -> Browse to start-server.bat')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Firewall Configuration:')
run.bold = True
doc.add_paragraph('Open Windows Defender Firewall -> Advanced Settings -> Inbound Rules -> New Rule')
doc.add_paragraph('Rule Type: Port -> TCP -> Specific local ports: 3000')
doc.add_paragraph('Action: Allow the connection')
doc.add_paragraph('Profile: Check Domain, Private, Public')
doc.add_paragraph('Name: Remote Monitor Server')

doc.add_page_break()

doc.add_heading('Type 2: Render.com (Cloud Server)', level=2)
doc.add_paragraph('Best for: Global access, no infrastructure management')
doc.add_paragraph('Cost: Free tier available (sleeps after 15 min inactivity)')
doc.add_paragraph('Access: Anywhere with internet')

p = doc.add_paragraph()
run = p.add_run('Setup Steps:')
run.bold = True

steps = [
    'Create account at https://render.com',
    'Click "New +" -> "Web Service"',
    'Connect your GitHub repository (puniteswra-spec/PU)',
    'Configure:',
    '   - Name: remotemonitor (or any name)',
    '   - Region: Choose closest to your agents',
    '   - Branch: main',
    '   - Root Directory: RemoteMonitor-Merged (or wherever server.source.js is)',
    '   - Runtime: Node',
    '   - Build Command: npm install',
    '   - Start Command: node server.source.js',
    '   - Instance Type: Free',
    'Click "Create Web Service"',
    'Wait for deployment (2-5 minutes)',
    'Your dashboard URL will be: https://your-app-name.onrender.com',
    'Update GitHub registry: Edit urls.ini in your repo to add: wss://your-app-name.onrender.com'
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(str(i) + '. ' + s)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Note on Free Tier:')
run.bold = True
doc.add_paragraph('Render free tier services sleep after 15 minutes of inactivity. The first request after sleep takes 30-60 seconds to wake up. For production use, consider the $7/month plan.')

doc.add_page_break()

doc.add_heading('Type 3: Linux VPS (Cloud Server)', level=2)
doc.add_paragraph('Best for: Production, 24/7 uptime, full control')
doc.add_paragraph('Cost: $5-10/month (DigitalOcean, Linode, AWS, etc.)')
doc.add_paragraph('Access: Anywhere with internet')

p = doc.add_paragraph()
run = p.add_run('Setup Steps:')
run.bold = True

p = doc.add_paragraph()
run = p.add_run('Step 1: Create VPS')
run.bold = True
doc.add_paragraph('Go to DigitalOcean (https://digitalocean.com) or any VPS provider')
doc.add_paragraph('Create a Droplet: Ubuntu 22.04, $6/month (1GB RAM)')
doc.add_paragraph('Note the IP address (e.g., 123.45.67.89)')

p = doc.add_paragraph()
run = p.add_run('Step 2: Connect via SSH')
run.bold = True
p = doc.add_paragraph()
run = p.add_run('ssh root@123.45.67.89')
run.font.name = 'Consolas'
run.font.size = Pt(9)

p = doc.add_paragraph()
run = p.add_run('Step 3: Install Node.js')
run.bold = True
p = doc.add_paragraph()
run = p.add_run('curl -fsSL https://deb.nodesource.com/setup_18.x | sudo -E bash -\nsudo apt-get install -y nodejs')
run.font.name = 'Consolas'
run.font.size = Pt(9)

p = doc.add_paragraph()
run = p.add_run('Step 4: Setup Application')
run.bold = True
p = doc.add_paragraph()
run = p.add_run('mkdir -p /opt/remotemonitor\ncd /opt/remotemonitor')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph('Upload files via SCP from your PC:')
p = doc.add_paragraph()
run = p.add_run('scp server.source.js package.json dashboard.html root@123.45.67.89:/opt/remotemonitor/')
run.font.name = 'Consolas'
run.font.size = Pt(9)

p = doc.add_paragraph()
run = p.add_run('npm install --production')
run.font.name = 'Consolas'
run.font.size = Pt(9)

p = doc.add_paragraph()
run = p.add_run('Step 5: Install PM2 (process manager)')
run.bold = True
p = doc.add_paragraph()
run = p.add_run('sudo npm install -g pm2\npm2 start server.source.js --name remotemonitor\npm2 save\npm2 startup')
run.font.name = 'Consolas'
run.font.size = Pt(9)

p = doc.add_paragraph()
run = p.add_run('Step 6: Open Firewall')
run.bold = True
p = doc.add_paragraph()
run = p.add_run('sudo ufw allow 3000/tcp\nsudo ufw enable')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph('Dashboard URL: http://123.45.67.89:3000')
doc.add_paragraph('Agent URL: ws://123.45.67.89:3000')

doc.add_page_break()

doc.add_heading('Type 4: Tunnel (No Server Setup)', level=2)
doc.add_paragraph('Best for: Temporary access, testing, no server infrastructure')
doc.add_paragraph('Cost: Free')
doc.add_paragraph('Access: Anywhere with internet (through tunnel)')

doc.add_paragraph('If you have a local server but want global access without a VPS:')

p = doc.add_paragraph()
run = p.add_run('Option A: Cloudflare Tunnel (Recommended)')
run.bold = True
doc.add_paragraph('Install cloudflared: winget install Cloudflare.cloudflared')
p = doc.add_paragraph()
run = p.add_run('cloudflared tunnel --url http://localhost:3000')
run.font.name = 'Consolas'
run.font.size = Pt(9)
doc.add_paragraph('Gives you: https://random-name.trycloudflare.com')
doc.add_paragraph('Agent URL: wss://random-name.trycloudflare.com (replace https:// with wss://)')

p = doc.add_paragraph()
run = p.add_run('Option B: bore.pub')
run.bold = True
doc.add_paragraph('Download bore from https://github.com/ekzhang/bore')
p = doc.add_paragraph()
run = p.add_run('bore local 3000 --to bore.pub')
run.font.name = 'Consolas'
run.font.size = Pt(9)
doc.add_paragraph('Gives you: bore.pub:XXXXX')
doc.add_paragraph('Agent URL: ws://bore.pub:XXXXX')

doc.add_page_break()

doc.add_heading('How to Change Servers', level=1)

doc.add_heading('Method 1: GitHub Registry (Recommended - Updates ALL Agents)', level=2)
doc.add_paragraph('This method updates ALL agents automatically within 10 minutes.')

steps = [
    'Go to your GitHub repository: https://github.com/puniteswra-spec/PU',
    'Open the file: urls.ini',
    'Edit the file - add your NEW server URL at the TOP (first line = highest priority)',
    'Example urls.ini:',
]
for s in steps:
    doc.add_paragraph(s)

p = doc.add_paragraph()
run = p.add_run('wss://new-server.onrender.com\nws://192.168.1.100:3000\nwss://pu-k752.onrender.com')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph('6. Commit the changes')
doc.add_paragraph('7. All agents will fetch the new URLs within 10 minutes and connect to the new server')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Where to find the registry URL in code:')
run.bold = True
doc.add_paragraph('main.go, line 46:')
p = doc.add_paragraph()
run = p.add_run('GitHubRegistryURL  = "https://raw.githubusercontent.com/puniteswra-spec/PU/main/urls.ini"')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_page_break()

doc.add_heading('Method 2: Dashboard Wizard (Push to Connected Agents)', level=2)
doc.add_paragraph('This method instantly switches all CONNECTED agents.')

steps = [
    'Open dashboard (http://localhost:3000 or Render URL)',
    'Login as admin (click lock icon, password: puneet12)',
    'Click "Update Remote" button',
    'Select server type:',
    '   - Render.com: Uses wss://pu-k752.onrender.com',
    '   - Local Server: Uses ws://<current-host>:3000',
    '   - VPS Server: Enter custom URL (e.g., ws://123.45.67.89:3000)',
    'Click "Apply"',
    'All connected agents switch immediately'
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(str(i) + '. ' + s)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('API endpoint used:')
run.bold = True
doc.add_paragraph('POST /api/switch-server')
doc.add_paragraph('Header: x-server-url: <new-ws-url>')
doc.add_paragraph('Auth: Basic Auth (puneet:puneet12)')

doc.add_heading('Method 3: Per-Agent Config Panel', level=2)
doc.add_paragraph('This method changes server for ONE specific agent.')

steps = [
    'On the agent PC, open browser to: http://localhost:8181',
    'This opens the SystemHelper Config Panel',
    'Enter new server URL in the input field',
    'Click "Add & Reconnect"',
    'Agent will restart and connect to the new server'
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(str(i) + '. ' + s)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('What this does:')
run.bold = True
doc.add_paragraph('Saves the URL to %APPDATA%\\SystemHelper\\urls.ini')
doc.add_paragraph('This file is read BEFORE the default URLs, so it takes priority')

doc.add_heading('Method 4: Edit urls.ini Directly', level=2)
doc.add_paragraph('This method changes server for ONE specific agent (manual).')

steps = [
    'On the agent PC, navigate to: %APPDATA%\\SystemHelper\\',
    'Open urls.ini in Notepad',
    'Add your new server URL on the FIRST LINE',
    'Save the file',
    'Restart the agent (or wait for next reconnect)'
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(str(i) + '. ' + s)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('urls.ini format:')
run.bold = True
p = doc.add_paragraph()
run = p.add_run('# One URL per line\n# First URL = highest priority\nwss://new-server.onrender.com\nws://192.168.1.100:3000')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_page_break()

doc.add_heading('Quick Reference: Where to Update URLs', level=1)

table2 = doc.add_table(rows=9, cols=3)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
headers2 = ['Where', 'File/Location', 'What to Change']
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

ref_data = [
    ['GitHub Registry', 'GitHub repo: urls.ini', 'Add new wss:// URL at top'],
    ['Agent Default', 'main.go line 44', 'DefaultServerURL constant'],
    ['Agent Fallback', 'main.go line 45', 'DirectServerIP constant'],
    ['Agent Local', 'main.go line 52', 'serverUrls array'],
    ['Agent Config', '%APPDATA%\\SystemHelper\\urls.ini', 'Add new URL at top'],
    ['Agent Next to .exe', 'urls.ini (same folder as .exe)', 'Add new URL at top'],
    ['Dashboard', 'dashboard.html line 256', 'wizardApply() function'],
    ['Server Auth', 'server.source.js lines 13-14', 'AUTH_USER / AUTH_PASS']
]
for r, row_data in enumerate(ref_data, 1):
    for c, val in enumerate(row_data):
        table2.rows[r].cells[c].text = val
        for paragraph in table2.rows[r].cells[c].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_page_break()

doc.add_heading('Remote Assistant Link', level=1)
doc.add_paragraph('The Remote Assistant feature allows anyone to share their screen or camera from a web browser — no installation required. This works on Windows, Mac, Linux, iOS, and Android.')

doc.add_heading('Remote Assistant URL', level=2)
doc.add_paragraph('The link is: https://your-server.com/remote-assistant')
doc.add_paragraph('For your current server: https://pu-k752.onrender.com/remote-assistant')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('How it works:')
run.bold = True
doc.add_paragraph('1. Share the Remote Assistant link with the person who needs help', style='List Bullet')
doc.add_paragraph('2. They open it in any browser (Chrome, Safari, Edge, Firefox)', style='List Bullet')
doc.add_paragraph('3. On desktop: they click "Start Sharing" and select their screen/window/tab', style='List Bullet')
doc.add_paragraph('4. On mobile: they click "Share Camera" to share their camera view', style='List Bullet')
doc.add_paragraph('5. Their screen appears as a tile in your dashboard grid', style='List Bullet')
doc.add_paragraph('6. You can view their screen live — no control, but you can guide them verbally', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Where to find the link:')
run.bold = True
doc.add_paragraph('Dashboard topbar has a "Remote Assistant" button (green link) that opens the page in a new tab')
doc.add_paragraph('You can copy this URL and share it via email, chat, SMS, etc.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('If you change servers:')
run.bold = True
doc.add_paragraph('The Remote Assistant link automatically uses your current server URL. No need to update anything — just share the new server URL with /remote-assistant at the end.')

doc.add_page_break()

doc.add_heading('Migration Checklist', level=1)
doc.add_paragraph('When moving from one server to another, follow this checklist:')

checklist = [
    'Set up the new server (follow setup steps above)',
    'Test the new server locally (open dashboard in browser)',
    'Test Remote Assistant link: https://new-server.com/remote-assistant',
    'Update GitHub registry (urls.ini) with new URL',
    'Wait 10 minutes for agents to fetch new URLs',
    'Verify agents are connecting to new server (check dashboard)',
    'Update dashboard.html wizard URL if different from Render',
    'Update main.go constants if changing default server',
    'Test remote control on new server',
    'Test tunnel/expose on new server',
    'Update documentation with new server URL'
]
for i, item in enumerate(checklist, 1):
    doc.add_paragraph('[ ] ' + str(i) + '. ' + item)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Important Notes:')
run.bold = True

notes = [
    'Agents try URLs in order - put the new server FIRST in urls.ini',
    'The agent will keep trying until it connects - no manual restart needed',
    'If the old server is still running, agents may connect to it first',
    'Use the Dashboard Wizard for instant switching of connected agents',
    'GitHub registry is the best method for updating ALL agents at once',
    'Test the new server BEFORE switching to avoid downtime'
]
for note in notes:
    doc.add_paragraph('- ' + note, style='List Bullet')

doc.save('P:/Opencode/RemoteMonitor-Merged/Server_Change_Guide.docx')
print('Server change guide saved successfully')
